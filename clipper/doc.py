"""文档剪藏模块 - PDF / Word 正文抽取为 Markdown + 原文件存档

依赖：python-docx（.docx）、pypdf（.pdf）。均为纯 Python 无原生编译。
抽取失败 / 加密文档 → 仅存原文件 + 失败占位（由调用方 clip_file 处理 md 头）。
"""

import contextlib
import shutil
from collections.abc import Callable
from datetime import datetime
from pathlib import Path
from typing import Any

from clipper.logging import get_logger

_log = get_logger("clipper.doc")


def clip_doc(
    file_path: str,
    output_dir: Path,
    category: str = "其他收藏",
    source_url: str | None = None,
    category_fn: Callable[[str, str], str] | None = None,
) -> dict[str, Any]:
    """
    抽取 PDF/Word 正文为 Markdown 并存原文件到条目目录。

    Returns:
        dict: {success, title, md_file, files, page_count, word_count,
               doc_type, warnings, error}
    """
    result: dict[str, Any] = {
        "success": False,
        "title": "文档",
        "md_file": None,
        "files": [],
        "page_count": None,
        "word_count": None,
        "doc_type": None,
        "warnings": [],
        "error": None,
        "fetch_backend": "local",  # FR-012
    }

    src = Path(file_path)
    suffix = src.suffix.lower()
    if suffix == ".docx":
        doc_type = "docx"
    elif suffix == ".pdf":
        doc_type = "pdf"
    else:
        result["error"] = f"不支持的文档类型: {suffix}"
        return result
    result["doc_type"] = doc_type

    # FR-013b: 文件大小前置校验
    from clipper.validators import validate_file_size

    size_ok, size_err = validate_file_size(src)
    if not size_ok:
        result["warnings"].append(size_err)
        _log.warning("doc_size_exceeded", src=str(src))

    output_dir.mkdir(parents=True, exist_ok=True)

    # 原文件存档
    try:
        dest_doc = output_dir / src.name
        shutil.copy2(str(src), str(dest_doc))
        result["files"] = [str(dest_doc)]
    except Exception as e:
        result["error"] = f"保存原文件失败: {e}"
        _log.warning("doc_save_failed", src=str(src), error=str(e))
        return result

    # 抽取正文
    body_md = ""
    title = src.stem
    page_count = None
    word_count = None
    capture_status = "成功"

    try:
        if doc_type == "docx":
            from docx import Document

            doc = Document(str(src))
            lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    lines.append("")
                    continue
                style = (para.style.name if para.style is not None else "") or ""
                style = style.lower()
                if "heading 1" in style or "title" in style:
                    lines.append(f"# {text}")
                elif "heading 2" in style:
                    lines.append(f"## {text}")
                elif "heading 3" in style:
                    lines.append(f"### {text}")
                elif "heading" in style:
                    lines.append(f"#### {text}")
                else:
                    lines.append(text)
            body_md = "\n".join(lines)
            # 粗略字数
            word_count = sum(len(p.text) for p in doc.paragraphs if p.text.strip())
        elif doc_type == "pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(src))
            page_count = len(reader.pages)
            pages_text = []
            for i, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""
                pages_text.append(f"## 页 {i}\n\n{text}")
            body_md = "\n\n".join(pages_text)
            word_count = sum(len(t) for t in pages_text)
    except Exception as e:
        # 加密 / 抽取失败 → 失败占位（仍保留原文件）(FR-013)
        capture_status = "失败"
        result["warnings"].append(f"正文抽取失败: {e}，已保留原文件")
        result["error"] = f"正文抽取失败: {e}"
        body_md = ""
        _log.warning("doc_extract_failed", src=str(src), error=str(e)[:200])

    result["page_count"] = page_count
    result["word_count"] = word_count

    # 标题用文档名
    title = src.stem or "文档"
    result["title"] = title

    # 回算领域
    if category_fn is not None:
        with contextlib.suppress(Exception):
            category = category_fn(title, body_md[:500]) or category
    result["category"] = category

    # 构建 md
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header_lines = [
        f"# {title}",
        "",
        f"> 📅 剪藏时间：{now}",
        "> 🏷️ 来源类型：doc",
        f"> 📁 领域：{category}",
        f"> ✅ 抓取状态：{capture_status}",
        f"> 📑 文件类型：{doc_type}",
        (f"> 🔗 来源：[{source_url}]({source_url})" if source_url else "> 📎 原文件已存档"),
        "",
        "---",
        "",
    ]
    if capture_status == "成功":
        body_section = body_md if body_md.strip() else "> 正文为空"
        meta_section = [
            "---",
            "",
            "## 📋 元数据",
            "",
            "- **来源类型**：doc",
            f"- **文件类型**：{doc_type}",
            f"- **领域**：{category}",
            "- **抓取状态**：成功",
            (f"- **页数**：{page_count}" if page_count else ""),
            (f"- **字数**：{word_count}" if word_count else ""),
            f"- **原文件**：{src.name}",
            f"- **剪藏时间**：{now}",
        ]
    else:
        body_section = (
            "## ❗ 失败原因\n\n"
            f"{result.get('error', '正文抽取失败')}\n\n"
            "原文件已存档，可手动打开或交 clawbot 识图兜底核对。"
        )
        meta_section = [
            "---",
            "",
            "## 📋 元数据",
            "",
            "- **来源类型**：doc",
            f"- **文件类型**：{doc_type}",
            f"- **领域**：{category}",
            "- **抓取状态**：失败",
            f"- **原文件**：{src.name}",
            f"- **剪藏时间**：{now}",
        ]

    md_content = (
        "\n".join(header_lines) + body_section + "\n\n" + "\n".join(filter(None, meta_section))
    )

    md_path = output_dir / "doc.md"
    md_path.write_text(md_content, encoding="utf-8")
    result["md_file"] = str(md_path)

    # success 仅在成功抽取时置真；失败时仍落占位 md
    result["success"] = capture_status == "成功"
    if not result["success"]:
        # 仍标记有 md 产出（用于索引记录占位）
        result["md_file"] = str(md_path)
    return result
